import os
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.popup import Popup
import whisper
import openai
from docx import Document

API_KEY_FILE = "api_key.txt"  # File to store API key


class MeetingMinutesApp(App):
    def build(self):
        self.api_key = self.load_api_key()

        if self.api_key:
            openai.api_key = self.api_key
            return self.main_screen()
        else:
            return self.api_key_screen()

    def api_key_screen(self):
        """Screen for entering the API key"""
        layout = BoxLayout(orientation='vertical')
        layout.add_widget(Label(text="Enter OpenAI API Key:"))

        self.api_input = TextInput(hint_text="Paste your OpenAI API Key here", password=True)
        layout.add_widget(self.api_input)

        submit_button = Button(text="Save API Key")
        submit_button.bind(on_press=self.save_api_key)
        layout.add_widget(submit_button)

        return layout

    def save_api_key(self, instance):
        """Save API key to file and switch to main screen"""
        key = self.api_input.text.strip()
        if key.startswith("sk-"):
            with open(API_KEY_FILE, "w") as file:
                file.write(key)

            self.api_key = key
            openai.api_key = key
            self.root.clear_widgets()
            self.root.add_widget(self.main_screen())
        else:
            self.show_popup("Error", "Invalid API Key. Please enter a valid key.")

    def load_api_key(self):
        """Load API key from file"""
        if os.path.exists(API_KEY_FILE):
            with open(API_KEY_FILE, "r") as file:
                return file.read().strip()
        return None

    def main_screen(self):
        """Main app screen after API key is set"""
        layout = BoxLayout(orientation='vertical')
        layout.add_widget(Label(text="Select an audio file for transcription"))

        self.file_chooser = FileChooserIconView()
        layout.add_widget(self.file_chooser)

        self.transcribe_button = Button(text="Transcribe & Generate Minutes")
        self.transcribe_button.bind(on_press=self.transcribe_audio)
        layout.add_widget(self.transcribe_button)

        return layout

    def transcribe_audio(self, instance):
        """Handle transcription and meeting minutes generation"""
        selected_file = self.file_chooser.selection
        if not selected_file:
            self.show_popup("Error", "Please select an audio file.")
            return

        audio_path = selected_file[0]
        self.show_popup("Processing", "Transcribing... Please wait.")

        try:
            model = whisper.load_model("base")
            result = model.transcribe(audio_path)
            transcription = result["text"]

            summary = self.generate_meeting_minutes(transcription)
            self.save_as_docx(summary)
            self.show_popup("Success", "Meeting minutes saved as 'meeting_minutes.docx'")

        except Exception as e:
            self.show_popup("Error", str(e))

    def generate_meeting_minutes(self, transcription):
        """Use OpenAI to generate meeting minutes"""
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Summarize the following meeting transcript."},
                {"role": "user", "content": transcription}
            ]
        )
        return response["choices"][0]["message"]["content"]

    def save_as_docx(self, content):
        """Save meeting minutes as a Word document"""
        doc = Document()
        doc.add_heading("Meeting Minutes", level=1)
        doc.add_paragraph(content)
        doc.save("meeting_minutes.docx")

    def show_popup(self, title, message):
        """Display a popup message"""
        popup_layout = BoxLayout(orientation='vertical')
        popup_label = Label(text=message)
        close_button = Button(text="OK")
        popup_layout.add_widget(popup_label)
        popup_layout.add_widget(close_button)

        popup = Popup(title=title, content=popup_layout, size_hint=(0.6, 0.4))
        close_button.bind(on_press=popup.dismiss)
        popup.open()


if __name__ == "__main__":
    MeetingMinutesApp().run()
