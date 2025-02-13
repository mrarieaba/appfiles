from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.popup import Popup
from kivy.uix.progressbar import ProgressBar
import whisper
import openai
from docx import Document
import os

# Set your OpenAI API key
OPENAI_API_KEY = "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"
openai.api_key = OPENAI_API_KEY

class MeetingMinutesApp(App):
    def build(self):
        self.layout = BoxLayout(orientation='vertical')
        self.label = Label(text="Select an audio file for transcription")
        self.layout.add_widget(self.label)
        
        self.file_chooser = FileChooserIconView()
        self.layout.add_widget(self.file_chooser)
        
        self.transcribe_button = Button(text="Transcribe & Generate Minutes")
        self.transcribe_button.bind(on_press=self.transcribe_audio)
        self.layout.add_widget(self.transcribe_button)
        
        return self.layout
    
    def transcribe_audio(self, instance):
        selected_file = self.file_chooser.selection
        if not selected_file:
            self.show_popup("Error", "Please select an audio file.")
            return
        
        audio_path = selected_file[0]
        self.label.text = "Transcribing... Please wait."
        
        try:
            model = whisper.load_model("base")  # Small model for faster processing
            result = model.transcribe(audio_path)
            transcription = result["text"]
            self.label.text = "Generating meeting minutes..."
            
            summary = self.generate_meeting_minutes(transcription)
            self.save_as_docx(summary)
            self.show_popup("Success", "Meeting minutes saved as 'meeting_minutes.docx'")
            
        except Exception as e:
            self.show_popup("Error", str(e))
    
    def generate_meeting_minutes(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Summarize the following meeting transcript."},
                {"role": "user", "content": transcription}
            ]
        )
        return response["choices"][0]["message"]["content"]
    
    def save_as_docx(self, content):
        doc = Document()
        doc.add_heading("Meeting Minutes", level=1)
        doc.add_paragraph(content)
        doc.save("meeting_minutes.docx")
    
    def show_popup(self, title, message):
        popup_layout = BoxLayout(orientation='vertical')
        popup_label = Label(text=message)
        close_button = Button(text="OK")
        popup_layout.add_widget(popup_label)
        popup_layout.add_widget(close_button)
        
        popup = Popup(title=title, content=popup_layout, size_hint=(0.6, 0.4))
        close_button.bind(on_press=popup.dismiss)
        popup.open()

if __name__ == "__main__":
    MeetingMinutesApp().run()
