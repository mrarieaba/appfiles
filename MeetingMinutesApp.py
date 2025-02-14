from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.popup import Popup
from kivy.uix.textinput import TextInput
from kivy.storage.jsonstore import JsonStore
import openai
from docx import Document
import os

# Store API key securely
store = JsonStore("config.json")

class MeetingMinutesApp(App):
    def build(self):
        self.layout = BoxLayout(orientation="vertical")
        self.label = Label(text="Select an audio file for transcription")
        self.layout.add_widget(self.label)

        self.file_chooser = FileChooserIconView()
        self.layout.add_widget(self.file_chooser)

        self.transcribe_button = Button(text="Transcribe & Generate Minutes")
        self.transcribe_button.bind(on_press=self.transcribe_audio)
        self.layout.add_widget(self.transcribe_button)

        # Check for API key, prompt if missing
        if not store.exists("openai_api_key"):
            self.ask_api_key()

        return self.layout

    def ask_api_key(self):
        """Prompts the user to enter their OpenAI API key and saves it securely."""
        popup_layout = BoxLayout(orientation="vertical", spacing=10)
        self.api_input = TextInput(hint_text="Enter OpenAI API Key", password=True, multiline=False)
        save_button = Button(text="Save API Key")

        popup_layout.add_widget(self.api_input)
        popup_layout.add_widget(save_button)

        self.api_popup = Popup(title="API Key Required", content=popup_layout, size_hint=(0.8, 0.4))
        save_button.bind(on_press=self.save_api_key)
        self.api_popup.open()

    def save_api_key(self, instance):
        """Saves the API key to local storage."""
        api_key = self.api_input.text.strip()
        if api_key:
            store.put("openai_api_key", key=api_key)
            openai.api_key = api_key
            self.api_popup.dismiss()
        else:
            self.show_popup("Error", "API key cannot be empty.")

    def transcribe_audio(self, instance):
        """Handles the audio transcription process using OpenAI's Whisper API."""
        selected_file = self.file_chooser.selection
        if not selected_file:
            self.show_popup("Error", "Please select an audio file.")
            return

        # Ensure API key is set
        if store.exists("openai_api_key"):
            openai.api_key = store.get("openai_api_key")["key"]
        else:
            self.ask_api_key()
            return

        audio_path = selected_file[0]
        self.label.text = "Transcribing... Please wait."

        try:
            with open(audio_path, "rb") as audio_file:
                transcription = openai.Audio.transcribe(model="whisper-1", file=audio_file)

            self.label.text = "Generating meeting minutes..."
            summary = self.generate_meeting_minutes(transcription["text"])
            self.save_as_docx(summary)
            self.show_popup("Success", "Meeting minutes saved as 'meeting_minutes.docx'.")

        except Exception as e:
            self.show_popup("Error", str(e))

    def generate_meeting_minutes(self, transcription):
        """Uses OpenAI to generate a structured summary, key points, action items, and sentiment analysis."""
        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=0,
            messages=[
                {"role": "system", "content": "Provide a structured summary of the meeting transcript, including:\n1. Abstract Summary\n2. Key Points\n3. Action Items\n4. Sentiment Analysis."},
                {"role": "user", "content": transcription}
            ]
        )
        return response["choices"][0]["message"]["content"]

    def save_as_docx(self, content):
        """Saves the generated summary as a Word document."""
        doc = Document()
        doc.add_heading("Meeting Minutes", level=1)
        doc.add_paragraph(content)
        doc.save("meeting_minutes.docx")

    def show_popup(self, title, message):
        """Displays pop-up messages for user feedback."""
        popup_layout = BoxLayout(orientation="vertical")
        popup_label = Label(text=message)
        close_button = Button(text="OK")
        popup_layout.add_widget(popup_label)
        popup_layout.add_widget(close_button)

        popup = Popup(title=title, content=popup_layout, size_hint=(0.6, 0.4))
        close_button.bind(on_press=popup.dismiss)
        popup.open()


if __name__ == "__main__":
    MeetingMinutesApp().run()
